import calendar
import csv
import datetime
import json
from datetime import date, timedelta
from decimal import Decimal, InvalidOperation
from urllib.parse import quote

from django.contrib import messages
from django.db.models import Sum
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.utils.translation import gettext as _
from django.views.decorators.http import require_POST

from staff.decorators import role_required
from staff.models import StaffProfile

from .forms import AccountCategoryForm, BudgetForm, CashBookForm, CashHandoverForm
from .labels import localized_category_name
from .models import AccountCategory, Budget, CashBook, CashHandover
from .services import (
    CURRENCIES,
    STATEMENT_GROUPS,
    group_by_category,
    group_by_category_and_payment,
    grouped_report,
    operational_rows,
    payment_group,
    payment_group_label,
    pl_bucket,
    statement_opening_balance,
    statement_rows,
    totals_by_currency,
    totals_by_payment_group,
    unified_transactions,
)


manager_required = role_required(StaffProfile.Role.MANAGER)

CURRENCY_SYMBOLS = {"LAK": "₭", "THB": "฿", "USD": "$"}

# ລາຍງານແຍກຕາມຊ່ອງທາງເຄີຍແຍກເປັນສີ່ໜ້າ (ຮັບ/ຈ່າຍ × ສົດ/ໂອນ). ດຽວນີ້ລວມເປັນ
# ໃບແຈ້ງຍອດອັນດຽວຕໍ່ກະເປົ໋າ ທີ່ມີທັງລາຍຮັບ ແລະ ລາຍຈ່າຍ — ລິ້ງເກົ່າຍັງໃຊ້ໄດ້.
LEGACY_PAYMENT_REPORT_TYPES = {
    "in_cash": "cash",
    "out_cash": "cash",
    "in_transfer": "bank",
    "out_transfer": "bank",
    "bank_account": "bank",
}


def _parse_date(value, fallback):
    try:
        return date.fromisoformat(value)
    except (TypeError, ValueError):
        return fallback


@manager_required
def dashboard(request):
    today = timezone.localdate()
    selected_date = _parse_date(request.GET.get("date"), today)
    currency = request.GET.get("currency", "")
    transaction_type = request.GET.get("type", "")
    category_id = request.GET.get("category", "")
    query = request.GET.get("q", "").strip()

    daily_totals = totals_by_currency(selected_date, selected_date)
    all_totals = totals_by_currency()
    transactions = unified_transactions(
        selected_date,
        selected_date,
        currency,
        transaction_type,
        query,
        category_id,
    )

    month_start = selected_date.replace(day=1)
    monthly_totals = totals_by_currency(month_start, selected_date)
    budget_rows = []
    for budget in Budget.objects.select_related("category").filter(month=month_start):
        actual = (
            CashBook.objects.filter(
                status=CashBook.Status.CONFIRMED,
                transaction_type="OUT",
                category=budget.category,
                currency=budget.currency,
                date__range=(month_start, selected_date),
            ).aggregate(total=Sum("amount"))["total"]
            or Decimal("0")
        )
        percent = min(100, round((actual / budget.amount) * 100)) if budget.amount else 0
        budget_rows.append(
            {
                "budget": budget,
                "category_name": localized_category_name(budget.category),
                "actual": actual,
                "remaining": budget.amount - actual,
                "percent": percent,
            }
        )

    categories = [
        {"pk": category.pk, "name": localized_category_name(category)}
        for category in AccountCategory.objects.filter(is_active=True)
    ]

    # ການແຍກ ເງິນສົດ / ບັນຊີ ຄິດຈາກລາຍການທັງມື້ ບໍ່ແມ່ນຈາກລາຍການທີ່ຖືກກັ່ນຕອງ
    # ຢູ່ໜ້າຈໍ ເພື່ອບໍ່ໃຫ້ຕົວກັ່ນຕອງເຮັດໃຫ້ຍອດການ໌ດປ່ຽນຕາມ.
    day_rows = unified_transactions(selected_date, selected_date)
    for code, groups in _payment_breakdown(day_rows).items():
        daily_totals[code]["split"] = {
            group: {**stats, "balance": stats["income"] - stats["expense"]}
            for group, stats in groups.items()
        }

    return render(
        request,
        "accounting/dashboard.html",
        {
            "active_nav": "accounting",
            "selected_date": selected_date,
            "transactions": transactions,
            "daily_totals": daily_totals,
            "whatsapp_text": _whatsapp_text(
                _("Hug ເກີບ AI daily ledger summary"),
                f"{_('Date')}: {selected_date:%d/%m/%Y}",
                daily_totals,
            ),
            "all_totals": all_totals,
            "monthly_totals": monthly_totals,
            "currencies": CURRENCIES,
            "categories": categories,
            "filters": {
                "currency": currency,
                "type": transaction_type,
                "category": str(category_id),
                "q": query,
            },
            "budget_rows": budget_rows,
        },
    )


@manager_required
def transaction_create(request):
    form = CashBookForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        transaction = form.save(commit=False)
        transaction.created_by = request.user
        transaction.updated_by = request.user
        transaction.save()
        messages.success(request, _("Transaction saved successfully."))
        return redirect(f"{reverse('accounting:dashboard')}?date={transaction.date}")
    return render(
        request,
        "accounting/transaction_form.html",
        {"form": form, "active_nav": "accounting", "page_title": _("Add accounting transaction")},
    )


@manager_required
def transaction_edit(request, pk):
    transaction = get_object_or_404(CashBook, pk=pk)
    form = CashBookForm(request.POST or None, request.FILES or None, instance=transaction)
    if form.is_valid():
        transaction = form.save(commit=False)
        transaction.updated_by = request.user
        transaction.save()
        messages.success(request, _("Transaction updated successfully."))
        return redirect(f"{reverse('accounting:dashboard')}?date={transaction.date}")
    return render(
        request,
        "accounting/transaction_form.html",
        {"form": form, "transaction": transaction, "active_nav": "accounting", "page_title": _("Edit accounting transaction")},
    )


@manager_required
@require_POST
def transaction_delete(request, pk):
    transaction = get_object_or_404(CashBook, pk=pk)
    view_date = transaction.date
    transaction.delete()
    messages.success(request, _("Transaction deleted successfully."))
    return redirect(f"{reverse('accounting:dashboard')}?date={view_date}")


def _report_filters(request):
    today = timezone.localdate()
    default_start = today.replace(day=1)
    start_date = _parse_date(request.GET.get("start"), default_start)
    end_date = _parse_date(request.GET.get("end"), today)
    if start_date > end_date:
        start_date, end_date = end_date, start_date
    grouping = request.GET.get("group", "daily")
    if grouping not in {"daily", "monthly", "yearly"}:
        grouping = "daily"
    return start_date, end_date, grouping


@manager_required
def report(request):
    start_date, end_date, grouping = _report_filters(request)
    rows = grouped_report(start_date, end_date, grouping)
    totals = totals_by_currency(start_date, end_date)
    max_amount = max(
        [row["income"] for row in rows] + [row["expense"] for row in rows] + [Decimal("0")]
    )
    for row in rows:
        row["income_percent"] = round((row["income"] / max_amount) * 100) if max_amount else 0
        row["expense_percent"] = round((row["expense"] / max_amount) * 100) if max_amount else 0
    return render(
        request,
        "accounting/report.html",
        {
            "active_nav": "accounting",
            "start_date": start_date,
            "end_date": end_date,
            "grouping": grouping,
            "rows": rows,
            "totals": totals,
            "whatsapp_text": _whatsapp_text(
                # ຂຽນຕົວອັກສອນລາວກົງໆ — ໃຊ້ \\u escape ແລ້ວ msgid ທີ່ xgettext ດຶງໄດ້
                # ຈະບໍ່ກົງກັບຂໍ້ຄວາມທີ່ Python ສົ່ງມາຕອນແລ່ນ ເຮັດໃຫ້ແປບໍ່ຕິດຕະຫຼອດ
                _("Hug ເກີບ AI financial report"),
                f"{_('Period')}: {start_date:%d/%m/%Y} – {end_date:%d/%m/%Y}",
                totals,
            ),
        },
    )


def _report_export_rows(start_date, end_date, grouping):
    """ຫົວຕາຕະລາງ + ແຖວຂໍ້ມູນ ທີ່ທຸກຮູບແບບການສົ່ງອອກໃຊ້ຮ່ວມກັນ"""
    headers = [_("Period"), _("Currency"), _("Income"), _("Expense"), _("Net")]
    rows = [
        [row["period"], row["currency"], row["income"], row["expense"], row["balance"]]
        for row in grouped_report(start_date, end_date, grouping)
    ]
    return headers, rows


@manager_required
def export_report(request):
    """ສົ່ງອອກລາຍງານເປັນ CSV / Excel / Word / PDF

    Excel ແລະ Word ຕ້ອງການ openpyxl ແລະ python-docx. ຖ້າຍັງບໍ່ໄດ້ຕິດຕັ້ງ
    ຈະແຈ້ງເຕືອນ ແລ້ວພາກັບໄປໜ້າລາຍງານ ແທນທີ່ຈະລົ້ມທັງໜ້າ.
    """
    start_date, end_date, grouping = _report_filters(request)
    export_format = (request.GET.get("format") or "csv").lower().strip()
    headers, rows = _report_export_rows(start_date, end_date, grouping)
    stem = f"hug-kerb-accounting-{start_date}-{end_date}"
    back_url = (
        f"{reverse('accounting:report')}?start={start_date}&end={end_date}&group={grouping}"
    )

    if export_format == "excel":
        try:
            import openpyxl
            from openpyxl.styles import Alignment, Font
            from openpyxl.utils import get_column_letter
        except ImportError:
            messages.error(
                request,
                _("Excel export needs the openpyxl package. Download CSV instead."),
            )
            return redirect(back_url)

        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Accounting Report"
        sheet.append(headers)
        for cell in sheet[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")
        for row in rows:
            sheet.append(row)
        for index, width in enumerate((18, 12, 18, 18, 18), start=1):
            sheet.column_dimensions[get_column_letter(index)].width = width

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = f'attachment; filename="{stem}.xlsx"'
        workbook.save(response)
        return response

    if export_format == "word":
        try:
            from docx import Document
        except ImportError:
            messages.error(
                request,
                _("Word export needs the python-docx package. Download CSV instead."),
            )
            return redirect(back_url)

        document = Document()
        document.add_heading(f"Hug ເກີບ AI — Financial Report ({grouping})", 0)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for cell, title in zip(table.rows[0].cells, headers):
            cell.text = str(title)
        for row in rows:
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.text = f"{value:,}" if isinstance(value, Decimal) else str(value)

        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="{stem}.docx"'
        document.save(response)
        return response

    if export_format == "pdf":
        # ໃຊ້ຕົວຊ່ວຍ PDF ຂອງໂປຣເຈັກ (WeasyPrint + ຝັງ Noto Sans Lao) ແທນ
        # reportlab ເພື່ອໃຫ້ຫົວຕາຕະລາງພາສາລາວອອກມາອ່ານໄດ້ ບໍ່ເປັນສີ່ຫຼ່ຽມດຳ.
        from core.pdf_fonts import pdf_font_context, pdf_response

        html = render(
            request,
            "accounting/report_pdf.html",
            {
                "headers": headers,
                "rows": rows,
                "start_date": start_date,
                "end_date": end_date,
                "grouping": grouping,
                **pdf_font_context(),
            },
        ).content.decode("utf-8")
        response = pdf_response(html, f"{stem}.pdf")
        response["Content-Disposition"] = f'attachment; filename="{stem}.pdf"'
        return response

    response = HttpResponse(content_type="text/csv; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{stem}.csv"'
    response.write("﻿")
    writer = csv.writer(response)
    writer.writerow(headers)
    for row in rows:
        writer.writerow(row)
    return response


@manager_required
def export_daily_transactions(request):
    """ສົ່ງອອກລາຍການປະຈຳວັນເປັນ Excel

    ກັ່ນຕອງຕາມກະເປົ໋າໄດ້ດ້ວຍ ?payment=cash|bank|other|all ເພື່ອໃຫ້ບັນຊີດຶງ
    ສະເພາະລາຍການທີ່ຜ່ານບັນຊີທະນາຄານໄປກະທົບຍອດກັບທະນາຄານໄດ້.
    """
    today = timezone.localdate()
    view_date = _parse_date(request.GET.get("date"), today)
    payment = (request.GET.get("payment") or "all").lower().strip()
    if payment not in STATEMENT_GROUPS and payment != "all":
        payment = "all"

    rows = [
        row
        for row in unified_transactions(view_date, view_date)
        if payment == "all" or payment_group(row["raw_method"]) == payment
    ]
    rows.sort(key=lambda row: row["sort_at"])
    payment_label = _("All payment methods") if payment == "all" else payment_group_label(payment)

    headers = [
        _("Time"),
        _("Source"),
        _("Description"),
        _("Category"),
        _("Currency"),
        _("Method"),
        _("Income (+)"),
        _("Expense (-)"),
        _("Recorded by"),
    ]

    def _row_values(row):
        entry = row["object"]
        recorded_by = getattr(entry, "created_by", None)
        return [
            row["time"].strftime("%H:%M") if row["time"] else "",
            row["source_label"],
            row["description"],
            row["category"] or "-",
            row["currency"],
            row["method"],
            row["amount"] if row["type"] == "IN" else None,
            row["amount"] if row["type"] == "OUT" else None,
            str(recorded_by) if recorded_by else "-",
        ]

    totals = {}
    for row in rows:
        bucket = totals.setdefault(
            row["currency"], {"income": Decimal("0"), "expense": Decimal("0")}
        )
        bucket["income" if row["type"] == "IN" else "expense"] += row["amount"]

    stem = f"hug-kerb-cashbook-{payment}-{view_date:%Y%m%d}"

    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        # ບໍ່ມີ openpyxl ກໍ່ຍັງໄດ້ໄຟລ໌ — ຖອຍໄປ CSV ແທນທີ່ຈະບໍ່ໄດ້ຫຍັງເລີຍ
        messages.warning(
            request,
            _("Excel export needs the openpyxl package — exported as CSV instead."),
        )
        response = HttpResponse(content_type="text/csv; charset=utf-8")
        response["Content-Disposition"] = f'attachment; filename="{stem}.csv"'
        response.write("﻿")
        writer = csv.writer(response)
        writer.writerow(headers)
        for row in rows:
            writer.writerow(["" if value is None else value for value in _row_values(row)])
        return response

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = payment[:31]

    sheet.append([f"Hug ເກີບ AI — {_('Daily transaction ledger')} ({payment_label})"])
    sheet.append([f"{_('Date')}: {view_date:%d/%m/%Y}"])
    sheet.append([])
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    sheet["A1"].font = Font(bold=True, size=14)
    sheet["A2"].font = Font(bold=True, size=11, color="FF475569")

    header_row = 4
    sheet.append(headers)
    thin = Side(style="thin", color="FFE2E8F0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for cell in sheet[header_row]:
        cell.font = Font(bold=True, color="FFFFFFFF")
        cell.fill = PatternFill("solid", fgColor="FF0891B2")
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    money_format = "#,##0.00"
    for row in rows:
        sheet.append(_row_values(row))

    for sheet_row in sheet.iter_rows(
        min_row=header_row + 1, max_row=sheet.max_row, max_col=len(headers)
    ):
        for cell in sheet_row:
            cell.border = border
        sheet_row[0].alignment = Alignment(horizontal="center")
        sheet_row[4].alignment = Alignment(horizontal="center")
        sheet_row[5].alignment = Alignment(horizontal="center")
        sheet_row[6].number_format = money_format
        sheet_row[7].number_format = money_format

    # ── ບລັອກສະຫຼູບ ໜຶ່ງແຖວຕໍ່ໜຶ່ງສະກຸນເງິນ ──
    sheet.append([])
    summary_head = sheet.max_row + 1
    sheet.append(
        [
            _("Summary"), "", "", "", "",
            _("Currency"),
            _("Total income"),
            _("Total expense"),
            _("Net balance"),
        ]
    )
    for cell in sheet[summary_head]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="FFF1F5F9")

    for currency in CURRENCIES:
        if currency not in totals:
            continue
        income = totals[currency]["income"]
        expense = totals[currency]["expense"]
        sheet.append(["", "", "", "", "", currency, income, expense, income - expense])
        for column in (7, 8, 9):
            cell = sheet.cell(row=sheet.max_row, column=column)
            cell.number_format = money_format
            cell.font = Font(bold=True)

    for index, width in enumerate((10, 20, 60, 32, 12, 16, 18, 18, 22), start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = sheet.cell(row=header_row + 1, column=1)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="{stem}.xlsx"'
    workbook.save(response)
    return response


@manager_required
def category_detail(request, pk):
    category = get_object_or_404(AccountCategory, pk=pk)
    today = timezone.localdate()
    start_date = _parse_date(request.GET.get("start"), today.replace(day=1))
    end_date = _parse_date(request.GET.get("end"), today)
    transactions = CashBook.objects.filter(
        category=category, date__range=(start_date, end_date)
    ).select_related("created_by")
    total = transactions.filter(status=CashBook.Status.CONFIRMED).aggregate(total=Sum("amount"))["total"] or 0
    return render(
        request,
        "accounting/category_detail.html",
        {
            "active_nav": "accounting",
            "category": category,
            "transactions": transactions,
            "total": total,
            "start_date": start_date,
            "end_date": end_date,
        },
    )


@manager_required
def cash_handover(request):
    today = timezone.localdate()
    selected_date = _parse_date(request.GET.get("date") or request.POST.get("date"), today)
    selected_currency = request.GET.get("currency") or request.POST.get("currency") or "LAK"
    existing = CashHandover.objects.filter(
        date=selected_date, currency=selected_currency
    ).first()
    form = CashHandoverForm(
        request.POST or None,
        instance=existing,
        initial={"date": selected_date, "currency": selected_currency},
    )
    cash_totals = totals_by_currency(
        selected_date, selected_date, payment_method=CashBook.PaymentMethod.CASH
    )
    if form.is_valid():
        handover = form.save(commit=False)
        opening = handover.opening_balance or Decimal("0")
        handover.expected_amount = opening + cash_totals[handover.currency]["balance"]
        handover.handed_by = request.user
        handover.save()
        messages.success(request, _("Handover saved successfully."))
        return redirect(
            f"{reverse('accounting:cash_handover')}?date={handover.date}&currency={handover.currency}"
        )
    history = CashHandover.objects.select_related("handed_by")[:30]
    return render(
        request,
        "accounting/cash_handover.html",
        {
            "active_nav": "accounting",
            "form": form,
            "existing": existing,
            "cash_totals": cash_totals,
            "selected_date": selected_date,
            "selected_currency": selected_currency,
            "history": history,
            "currencies": CURRENCIES,
        },
    )


def _month_names():
    """ຊື່ເດືອນ — Django ບໍ່ມີ locale ລາວມາໃຫ້ ຈຶ່ງແປເອງຜ່ານ catalog ຂອງໂຄງການ"""
    return [
        _("January"), _("February"), _("March"), _("April"),
        _("May"), _("June"), _("July"), _("August"),
        _("September"), _("October"), _("November"), _("December"),
    ]


@manager_required
def yearly_cash_handover(request):
    """ໃບສົ່ງມອບເງິນປະຈຳປີ — ລວມການສົ່ງມອບເງິນລາຍວັນທັງປີໄວ້ໃນໜ້າດຽວ

    ໃຊ້ຢືນຢັນກັບເຈົ້າຂອງຮ້ານທ້າຍປີວ່າ ເງິນທີ່ນັບໄດ້ຕະຫຼອດປີ ກົງກັບຍອດຕາມລະບົບບໍ່
    ແລະ ເດືອນໃດມີວັນທີ່ຍອດຂາດ/ເກີນ ຈະຖືກຍົກຂຶ້ນມາໃຫ້ເຫັນທັນທີ.
    """
    today = timezone.localdate()
    year_str = request.GET.get("year", "")
    year = int(year_str) if year_str.isdigit() else today.year
    currency = request.GET.get("currency", "LAK")
    if currency not in CURRENCIES:
        currency = "LAK"

    handovers = list(
        CashHandover.objects.filter(date__year=year, currency=currency)
        .select_related("handed_by")
        .order_by("date")
    )

    month_names = _month_names()
    months = []
    year_expected = year_total_counted = year_difference = Decimal("0")
    for month in range(1, 13):
        month_rows = [item for item in handovers if item.date.month == month]
        expected = sum((item.expected_amount for item in month_rows), Decimal("0"))
        counted = sum((item.counted_amount for item in month_rows), Decimal("0"))
        difference = counted - expected
        year_expected += expected
        year_total_counted += counted
        year_difference += difference
        months.append(
            {
                "month": month,
                "label": month_names[month - 1],
                "days": len(month_rows),
                "expected": expected,
                "counted": counted,
                "difference": difference,
                # ວັນທີ່ນັບເງິນບໍ່ກົງກັບລະບົບ — ຈຸດທີ່ຕ້ອງອະທິບາຍໃຫ້ໄດ້
                "mismatch_days": sum(1 for item in month_rows if item.difference != 0),
            }
        )

    mismatches = [item for item in handovers if item.difference != 0]
    recorded_days = len(handovers)

    return render(
        request,
        "accounting/yearly_cash_handover.html",
        {
            "active_nav": "accounting",
            "year": year,
            "years": range(today.year, today.year - 6, -1),
            "currency": currency,
            "currencies": CURRENCIES,
            "months": months,
            "year_expected": year_expected,
            "year_counted": year_total_counted,
            "year_difference": year_difference,
            "recorded_days": recorded_days,
            "mismatches": mismatches,
            "first_handover": handovers[0] if handovers else None,
            "last_handover": handovers[-1] if handovers else None,
            "today": today,
        },
    )


@manager_required
def settings_view(request):
    today = timezone.localdate()
    category_form = AccountCategoryForm(prefix="category")
    budget_form = BudgetForm(
        prefix="budget", initial={"month": today.replace(day=1)}
    )

    if request.method == "POST":
        action = request.POST.get("action")
        if action == "add_category":
            category_form = AccountCategoryForm(request.POST, prefix="category")
            if category_form.is_valid():
                category_form.save()
                messages.success(request, _("Category added successfully."))
                return redirect("accounting:settings")
        elif action == "save_budget":
            budget_form = BudgetForm(request.POST, prefix="budget")
            if budget_form.is_valid():
                data = budget_form.cleaned_data
                Budget.objects.update_or_create(
                    month=data["month"].replace(day=1),
                    category=data["category"],
                    currency=data["currency"],
                    defaults={"amount": data["amount"]},
                )
                messages.success(request, _("Budget saved successfully."))
                return redirect("accounting:settings")

    return render(
        request,
        "accounting/settings.html",
        {
            "active_nav": "accounting",
            "category_form": category_form,
            "budget_form": budget_form,
            "income_categories": [
                {
                    "pk": category.pk,
                    "name": localized_category_name(category),
                    "color": category.color,
                    "is_active": category.is_active,
                }
                for category in AccountCategory.objects.filter(transaction_type="IN")
            ],
            "expense_categories": [
                {
                    "pk": category.pk,
                    "name": localized_category_name(category),
                    "color": category.color,
                    "is_active": category.is_active,
                }
                for category in AccountCategory.objects.filter(transaction_type="OUT")
            ],
            "budgets": [
                {
                    "pk": budget.pk,
                    "month": budget.month,
                    "category_name": localized_category_name(budget.category),
                    "currency": budget.currency,
                    "amount": budget.amount,
                }
                for budget in Budget.objects.select_related("category")[:36]
            ],
        },
    )


@manager_required
@require_POST
def category_toggle(request, pk):
    category = get_object_or_404(AccountCategory, pk=pk)
    category.is_active = not category.is_active
    category.save(update_fields=["is_active"])
    messages.success(request, _("Category status updated successfully."))
    return redirect("accounting:settings")


@manager_required
@require_POST
def budget_delete(request, pk):
    get_object_or_404(Budget, pk=pk).delete()
    messages.success(request, _("Budget deleted successfully."))
    return redirect("accounting:settings")


def _get_accounting_cycle_dates(month_val, year_val, today):
    """
    Returns (current_month, current_year, start_date, end_date) for the standard
    calendar month: 1st day to the last day of the month.
    """
    m = int(month_val) if (month_val and str(month_val).isdigit()) else None
    y = int(str(year_val).replace(',', '').strip()) if (year_val and str(year_val).replace(',', '').strip().isdigit()) else None

    if m is None or y is None:
        m = today.month
        y = today.year

    start_date = datetime.date(y, m, 1)
    _, num_days = calendar.monthrange(y, m)
    end_date = datetime.date(y, m, num_days)
    return m, y, start_date, end_date


def _group_by_category(transactions, transaction_type, currency_code):
    """ລວມຍອດຕາມໝວດ ແລ້ວແປງເປັນ float ໃຫ້ຝັ່ງ JSON ຂອງໜ້າສະຫຼູບໃຊ້ໄດ້"""
    return [
        {"description": row["description"], "amount": float(row["amount"])}
        for row in group_by_category(transactions, transaction_type, currency_code)
    ]


def _summary_income(transactions, currency_code):
    """ຍອດລາຍຮັບຂອງໃບສະຫຼຸບ — ຕ້ອງໃຊ້ກົດດຽວກັບການ໌ດຢູ່ໜ້າ dashboard

    ຄິດຜ່ານ pl_bucket() ຈຶ່ງໄດ້ຄ່າດຽວກັນກັບ totals_by_currency() ສະເໝີ
    (ຄືນເງິນຫັກອອກ, ການຍ້າຍເງິນພາຍໃນຖືກຕັດອອກກ່ອນສົ່ງເຂົ້າມາແລ້ວ)
    """
    total = Decimal("0")
    for row in transactions:
        if row["currency"] != currency_code:
            continue
        bucket, amount = pl_bucket(row)
        if bucket == "income":
            total += amount
    return total


def _payment_breakdown(transactions):
    """ແຍກລາຍຮັບ/ລາຍຈ່າຍຕາມກະເປົ໋າ ຕໍ່ແຕ່ລະສະກຸນເງິນ (ໃຊ້ໃນໜ້າສະຫຼູບ ແລະ WhatsApp)"""
    breakdown = {}
    for currency in CURRENCIES:
        rows = [
            row
            for row in transactions
            if row["currency"] == currency and row["status"] == "confirmed"
        ]
        totals = totals_by_payment_group(rows)
        breakdown[currency] = {
            group: {
                "income": float(totals["IN"][group]),
                "expense": float(totals["OUT"][group]),
            }
            for group in STATEMENT_GROUPS
        }
    return breakdown


def _whatsapp_text(title, period_label, totals):
    """ຂໍ້ຄວາມສະຫຼູບພ້ອມແຊເຂົ້າ WhatsApp — encode ໄວ້ໃຫ້ວາງໃນ wa.me ໄດ້ເລີຍ"""
    lines = [f"*📊 {title}*", f"📅 {period_label}", "━━━━━━━━━━━━━━━", ""]
    has_activity = False
    for currency in CURRENCIES:
        stats = totals.get(currency)
        if not stats or (not stats["income"] and not stats["expense"]):
            continue
        has_activity = True
        symbol = CURRENCY_SYMBOLS.get(currency, "")
        icon = "✅" if stats["balance"] >= 0 else "❌"
        lines.append(f"*💰 {currency} ({symbol})*")
        lines.append(f"   🟢 {_('Income')}: *{stats['income']:,.0f}* {symbol}")
        lines.append(f"   🔴 {_('Expense')}: *{stats['expense']:,.0f}* {symbol}")
        lines.append(f"   {icon} {_('Net balance')}: *{stats['balance']:,.0f}* {symbol}")
        lines.append("")
    if not has_activity:
        lines.append(f"❌ _{_('No transactions in this period')}_")
        lines.append("")
    lines.append("━━━━━━━━━━━━━━━")
    lines.append("👟 *Hug ເກີບ AI*")
    return quote("\n".join(lines))


@manager_required
def monthly_summary_financial(request):
    """
    Generate the Monthly Financial Summary (ໃບສະຫຼຸບການເງິນປະຈຳເດືອນ).
    Auto-populates data from unified transactions (including POS) for LAK, THB, and USD.
    """
    today = timezone.localdate()

    month_str = request.GET.get('month')
    year_str = request.GET.get('year')
    run_count = request.GET.get('run_count', '1')
    date_str = request.GET.get('date')
    currency = request.GET.get('currency', 'ALL')
    
    current_month, current_year, start_date, end_date = _get_accounting_cycle_dates(
        month_str, year_str, today
    )
    voucher_date = datetime.datetime.strptime(date_str, '%Y-%m-%d').date() if date_str else today

    # ໃບສະຫຼຸບເປັນເອກະສານກຳໄລ-ຂາດທຶນ ຈຶ່ງຕັດການຍ້າຍເງິນລະຫວ່າງກະເປົ໋າຂອງຮ້ານເອງ
    # ແລະ ແຖວທີ່ຍັງບໍ່ຢືນຢັນອອກ — ຄືກັນກັບການ໌ດຢູ່ໜ້າ dashboard ແລະ ໜ້າລາຍງານ
    transactions = operational_rows(unified_transactions(start_date, end_date))

    # Calculate incomes
    cash_in_lak = _summary_income(transactions, "LAK")
    cash_in_thb = _summary_income(transactions, "THB")
    cash_in_usd = _summary_income(transactions, "USD")

    # Build list of items for manual adjustment sheet editor
    default_items = []
    for t in transactions:
        bucket, amount = pl_bucket(t)
        if bucket == "expense":
            default_items.append({
                'description': t["description"],
                'currency': t["currency"],
                'amount': float(amount)
            })

    # Group by category per currency
    report_data = {}
    for cur in CURRENCIES:
        report_data[cur] = {
            'income': _group_by_category(transactions, 'IN', cur),
            'expense': _group_by_category(transactions, 'OUT', cur),
        }

    context = {
        'active_nav': 'accounting',
        'payment_breakdown_json': json.dumps(_payment_breakdown(transactions), ensure_ascii=False),
        'selected_month': f"{current_month:02d}",
        'selected_year': str(current_year),
        'run_count': run_count,
        'voucher_date': voucher_date.strftime('%Y-%m-%d'),
        'voucher_date_formatted': voucher_date.strftime('%d-%m-%Y'),
        'default_cash_lak': float(cash_in_lak),
        'default_cash_thb': float(cash_in_thb),
        'default_cash_usd': float(cash_in_usd),
        'default_items_json': json.dumps(default_items),
        'report_json': json.dumps(report_data, ensure_ascii=False),
        'selected_currency': currency,
        'back_url': '/accounting/report/',
    }
    return render(request, 'accounting/monthly_summary_financial.html', context)


@manager_required
def yearly_summary_financial(request):
    """
    Generate the Yearly Financial Summary (ໃບສະຫຼຸບການເງິນປະຈຳປີ).
    Auto-populates data from unified transactions grouped by category and currency for the selected year.
    """
    today = timezone.localdate()

    year_str = request.GET.get('year')
    run_count = request.GET.get('run_count', '1')
    date_str = request.GET.get('date')
    currency = request.GET.get('currency', 'ALL')
    
    current_year = int(year_str) if (year_str and year_str.isdigit()) else today.year
    voucher_date = datetime.datetime.strptime(date_str, '%Y-%m-%d').date() if date_str else today

    start_date = datetime.date(current_year, 1, 1)
    end_date = datetime.date(current_year, 12, 31)

    # ຄືກັບໃບສະຫຼຸບເດືອນ: ຕັດການຍ້າຍເງິນພາຍໃນ ແລະ ແຖວທີ່ຍັງບໍ່ຢືນຢັນອອກ
    transactions = operational_rows(unified_transactions(start_date, end_date))

    # Calculate incomes
    cash_in_lak = _summary_income(transactions, "LAK")
    cash_in_thb = _summary_income(transactions, "THB")
    cash_in_usd = _summary_income(transactions, "USD")

    # Retrieve expenses grouped by category and currency
    # Using python grouping from unified transactions
    from collections import defaultdict
    expense_groups = defaultdict(float)
    for t in transactions:
        bucket, amount = pl_bucket(t)
        if bucket == "expense":
            key = (t["category"], t["currency"])
            expense_groups[key] += float(amount)

    default_items = []
    for (category_name, cur), total_amount in sorted(expense_groups.items(), key=lambda x: -x[1]):
        default_items.append({
            'description': category_name or 'ລາຍຈ່າຍອື່ນໆ (Other Expenses)',
            'currency': cur,
            'amount': total_amount
        })
        
    if not default_items:
        default_items.append({
            'description': 'ລາຍຈ່າຍອື່ນໆ (Other Expenses)',
            'currency': 'LAK',
            'amount': 0.0
        })
        
    # Group by category per currency for P&L structure
    report_data = {}
    for cur in CURRENCIES:
        report_data[cur] = {
            'income': _group_by_category(transactions, 'IN', cur),
            'expense': _group_by_category(transactions, 'OUT', cur),
        }

    context = {
        'active_nav': 'accounting',
        'payment_breakdown_json': json.dumps(_payment_breakdown(transactions), ensure_ascii=False),
        'selected_year': str(current_year),
        'run_count': run_count,
        'voucher_date': voucher_date.strftime('%Y-%m-%d'),
        'voucher_date_formatted': voucher_date.strftime('%d-%m-%Y'),
        'default_cash_lak': float(cash_in_lak),
        'default_cash_thb': float(cash_in_thb),
        'default_cash_usd': float(cash_in_usd),
        'default_items_json': json.dumps(default_items),
        'report_json': json.dumps(report_data, ensure_ascii=False),
        'selected_currency': currency,
        'back_url': '/accounting/report/',
    }
    return render(request, 'accounting/yearly_summary_financial.html', context)


@manager_required
def category_detail_print(request):
    """
    Printable line-by-line breakdown of a single category for a month.
    Supports key='office' (ຄ່າໃຊ້ຈ່າຍຫ້ອງການ), key='other' (ລາຍຈ່າຍອື່ນໆ),
    key='other_income' (ລາຍຮັບອື່ນໆ), key='office_income' (ເງິນບໍລິຫານຫ້ອງການ),
    key='bank_income' (ລາຍຮັບບັນຊີທະນາຄານ) ແລະ key='bank_expense'
    (ລາຍຈ່າຍບັນຊີທະນາຄານ) — ສອງອັນທ້າຍລວມທຸກລາຍການທີ່ຜ່ານບັນຊີທະນາຄານ
    ບໍ່ວ່າຈະຢູ່ໝວດໃດ.
    """
    today = timezone.localdate()

    key = request.GET.get('key', 'office')
    currency = request.GET.get('currency', 'LAK')
    if currency not in CURRENCIES:
        currency = 'LAK'

    month_str = request.GET.get('month')
    year_str = request.GET.get('year')
    current_month, current_year, start_date, end_date = _get_accounting_cycle_dates(
        month_str, year_str, today
    )

    # Fetch unified transactions
    transactions = unified_transactions(start_date, end_date, currency=currency)

    # ບັນຊີທະນາຄານ = ທຸກລາຍການທີ່ຜ່ານກະເປົ໋າ "bank" ບໍ່ຈຳກັດໝວດ
    group_filter = None

    # Filter by category matching the key
    if key == 'other_income':
        match_fn = lambda c: 'ລາຍຮັບອື່ນ' in c or 'other income' in c.lower()
        title_lo = 'ລາຍຮັບອື່ນໆ (Other Income)'
        title_en = 'Other Income'
        transaction_type = 'IN'
    elif key == 'office_income':
        match_fn = lambda c: 'ບໍລິຫານ' in c or 'office administration' in c.lower()
        title_lo = 'ເງິນບໍລິຫານຫ້ອງການ (Office Administration)'
        title_en = 'Office Administration Income'
        transaction_type = 'IN'
    elif key == 'bank_income':
        match_fn = lambda c: True
        group_filter = 'bank'
        title_lo = 'ລາຍຮັບຜ່ານບັນຊີທະນາຄານ (Bank Income)'
        title_en = 'Bank Income'
        transaction_type = 'IN'
    elif key == 'bank_expense':
        match_fn = lambda c: True
        group_filter = 'bank'
        title_lo = 'ລາຍຈ່າຍຜ່ານບັນຊີທະນາຄານ (Bank Expenses)'
        title_en = 'Bank Expenses'
        transaction_type = 'OUT'
    elif key == 'other':
        match_fn = lambda c: 'ອື່ນໆ' in c or 'other expenses' in c.lower()
        title_lo = 'ລາຍຈ່າຍອື່ນໆ (Other Expenses)'
        title_en = 'Other Expenses'
        transaction_type = 'OUT'
    else:
        key = 'office'
        match_fn = lambda c: 'ຫ້ອງການ' in c or 'office expenses' in c.lower()
        title_lo = 'ຄ່າໃຊ້ຈ່າຍຫ້ອງການ'
        title_en = 'Office Expenses'
        transaction_type = 'OUT'

    matched = [
        t
        for t in transactions
        if t["type"] == transaction_type
        and t["status"] == "confirmed"
        and (group_filter is None or payment_group(t["raw_method"]) == group_filter)
        and match_fn(t["category"])
    ]
    matched.sort(key=lambda t: t["sort_at"])

    # ── ຍອດເງິນບໍລິຫານຫ້ອງການ ───────────────────────────────────────────
    # ສະເພາະໜ້າ 'office': ຫັກລາຍຈ່າຍແຕ່ລະລາຍການອອກຈາກເງິນບໍລິຫານທີ່ຮັບມາ
    # ໃນເດືອນນັ້ນ ແລ້ວສະແດງຍອດຄົງເຫຼືອໃນຕາຕະລາງເລີຍ.
    show_balance = (key == 'office')
    opening_balance = Decimal('0')
    if show_balance:
        opening_balance = sum(
            (
                t["amount"]
                for t in transactions
                if t["type"] == "IN"
                and t["status"] == "confirmed"
                and ('ບໍລິຫານ' in t["category"] or 'office administration' in t["category"].lower())
            ),
            Decimal('0'),
        )
        # ອະນຸຍາດໃຫ້ກຳນົດຍອດຕັ້ງຕົ້ນເອງ ກໍລະນີຍັງບໍ່ໄດ້ບັນທຶກເງິນບໍລິຫານໃນສະໝຸດ
        raw_budget = (request.GET.get('budget') or '').replace(',', '').strip()
        if raw_budget:
            try:
                opening_balance = Decimal(raw_budget)
            except (InvalidOperation, ValueError):
                pass

    items = []
    total = Decimal('0')
    running_balance = opening_balance
    for t in matched:
        running_balance -= t["amount"]
        total += t["amount"]
        items.append({
            'date': t["date"],
            'description': t["description"],
            'category': t["category"] or '-',
            'method': t["method"],
            'amount': t["amount"],
            'balance': running_balance,
        })

    context = {
        'key': key,
        'title_lo': title_lo,
        'title_en': title_en,
        'items': items,
        'total': total,
        'show_balance': show_balance,
        'opening_balance': opening_balance,
        'remaining_balance': opening_balance - total,
        'currency': currency,
        'selected_month': f"{current_month:02d}",
        'selected_year': str(current_year),
        'today': today,
    }
    return render(request, 'accounting/category_detail_print.html', context)


@manager_required
def payment_method_report(request):
    """ໃບແຈ້ງຍອດປະຈຳເດືອນ ແຍກຕາມກະເປົ໋າທີ່ເງິນນອນຢູ່

    report_type:
    - 'all'  : ລວມທຸກຊ່ອງທາງ ເປັນຕາຕະລາງລາຍລະອຽດ
    - 'cash' : ເງິນສົດໃນກຳປັ່ນ ຈັດເປັນໃບແຈ້ງຍອດ (ຈ່າຍອອກ / ຮັບເຂົ້າ / ຍອດເຫຼືອ)
    - 'bank' : ບັນຊີທະນາຄານ (ເງິນໂອນ + ສະແກນ QR) ຮູບແບບດຽວກັນ

    ໃບແຈ້ງຍອດເປີດດ້ວຍຍອດຍົກມາຈາກກ່ອນງວດ ແລ້ວປິດດ້ວຍບລັອກລວມ
    (ຈຳນວນ ແລະ ຍອດ ຈ່າຍອອກ/ຮັບເຂົ້າ ພ້ອມຍອດເຫຼືອທ້າຍງວດ).
    """
    today = timezone.localdate()

    currency = request.GET.get("currency", "LAK")
    if currency not in CURRENCIES:
        currency = "LAK"

    current_month, current_year, start_date, end_date = _get_accounting_cycle_dates(
        request.GET.get("month"), request.GET.get("year"), today
    )

    report_type = (request.GET.get("report_type") or "all").lower().strip()
    report_type = LEGACY_PAYMENT_REPORT_TYPES.get(report_type, report_type)
    if report_type not in ("all", "cash", "bank"):
        report_type = "all"

    # None = ໜ້າລວມ; ນອກນັ້ນຄືກະເປົ໋າທີ່ໃບແຈ້ງຍອດນີ້ເວົ້າເຖິງ
    statement_group = report_type if report_type in STATEMENT_GROUPS else None

    period_rows = statement_rows(start_date, end_date, currency)

    # ການ໌ດສະຫຼູບ ແລະ ຕາຕະລາງແຍກໝວດ ເປັນເລື່ອງກຳໄລ-ຂາດທຶນ ຈຶ່ງຕັດການຍ້າຍເງິນ
    # ພາຍໃນອອກ. ຕົວໃບແຈ້ງຍອດຂ້າງລຸ່ມຍັງລວມມັນໄວ້ ເພາະເງິນຍ້າຍຈິງ.
    operational_rows = [row for row in period_rows if not row["internal"]]
    internal_moved = sum(
        (row["amount"] for row in period_rows if row["internal"] and row["type"] == "OUT"),
        Decimal("0"),
    )

    full_totals = totals_by_payment_group(operational_rows)

    items = []
    opening_balance = statement_opening_balance(start_date, currency, statement_group)
    running_balance = opening_balance
    total_debit = Decimal("0")
    total_credit = Decimal("0")
    debit_count = 0
    credit_count = 0

    ledger_rows = (
        period_rows
        if statement_group is None
        else [row for row in period_rows if payment_group(row["raw_method"]) == statement_group]
    )

    for row in ledger_rows:
        # ລາຍຮັບບວກເຂົ້າ ລາຍຈ່າຍຫັກອອກ — ການຍ້າຍເງິນພາຍໃນຈຶ່ງລົງຖືກຝັ່ງເອງ
        # ໂດຍບໍ່ຕ້ອງມີກົດພິເສດ: ຂາອອກຈາກກຳປັ່ນເປັນ debit, ຂາເຂົ້າບັນຊີເປັນ credit.
        signed = row["amount"] if row["type"] == "IN" else -row["amount"]
        debit = -signed if signed < 0 else Decimal("0")
        credit = signed if signed > 0 else Decimal("0")
        running_balance += signed
        if debit:
            total_debit += debit
            debit_count += 1
        if credit:
            total_credit += credit
            credit_count += 1
        items.append(
            {
                "date": row["date"],
                "time": row["time"],
                "description": row["description"],
                "category": row["category"],
                "transaction_type": row["type"],
                "method": row["method"],
                "payment_group": payment_group(row["raw_method"]),
                "internal": row["internal"],
                "amount": row["amount"],
                "debit": debit,
                "credit": credit,
                "balance": running_balance,
            }
        )

    # ຕາຕະລາງແຍກໝວດຕ້ອງສ້າງຈາກກະເປົ໋າດຽວກັນກັບທີ່ລາຍງານເວົ້າເຖິງ ບໍ່ດັ່ງນັ້ນ
    # ໜ້າ 'bank' ຈະຂຶ້ນໝວດທີ່ບໍ່ເຄີຍຈ່າຍຜ່ານບັນຊີເລີຍ.
    breakdown_rows = (
        operational_rows
        if statement_group is None
        else [
            row
            for row in operational_rows
            if payment_group(row["raw_method"]) == statement_group
        ]
    )
    category_breakdown = {
        "IN": group_by_category_and_payment(
            [row for row in breakdown_rows if row["type"] == "IN"]
        ),
        "OUT": group_by_category_and_payment(
            [row for row in breakdown_rows if row["type"] == "OUT"]
        ),
    }

    report_titles = {
        "all": (
            "Payment Method Report (All Methods)",
            "ລາຍງານແຍກຕາມຊ່ອງທາງຊຳລະ (ທຸກຊ່ອງທາງ)",
        ),
        "cash": (
            "Payment Report: Income — Expense (Cash)",
            "ລາຍງານແຍກຕາມຊ່ອງທາງ: ລາຍຮັບ — ລາຍຈ່າຍ (ເງິນສົດ)",
        ),
        "bank": (
            "Payment Report: Income — Expense (Bank Account)",
            "ລາຍງານແຍກຕາມຊ່ອງທາງ: ລາຍຮັບ — ລາຍຈ່າຍ (ບັນຊີທະນາຄານ)",
        ),
    }
    title_en, title_lo = report_titles[report_type]

    grand_total_income = sum(full_totals["IN"].values(), Decimal("0"))
    grand_total_expense = sum(full_totals["OUT"].values(), Decimal("0"))

    return render(
        request,
        "accounting/payment_method_report.html",
        {
            "active_nav": "accounting",
            "items": items,
            "full_totals": full_totals,
            "category_breakdown": category_breakdown,
            "grand_total_income": grand_total_income,
            "grand_total_expense": grand_total_expense,
            "currency": currency,
            "report_type": report_type,
            "report_title_en": title_en,
            "report_title_lo": title_lo,
            "selected_month": f"{current_month:02d}",
            "selected_year": str(current_year),
            "today": today,
            # ເງິນທີ່ຍ້າຍລະຫວ່າງກະເປົ໋າຂອງຮ້ານເອງ — ບໍ່ນັບເປັນລາຍຈ່າຍໃນຍອດຂ້າງເທິງ
            # ແຕ່ບອກໄວ້ໃຫ້ອະທິບາຍໄດ້ວ່າເງິນຫາຍໄປໃສ.
            "internal_moved": internal_moved,
            "statement_group": statement_group,
            "statement_group_label": (
                payment_group_label(statement_group) if statement_group else ""
            ),
            "opening_balance": opening_balance,
            "closing_balance": running_balance,
            "total_debit": total_debit,
            "total_credit": total_credit,
            "debit_count": debit_count,
            "credit_count": credit_count,
        },
    )
